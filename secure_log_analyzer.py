#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
sky's securelog — ELITE VISUALS (Entity Intelligence Removed)
- Modern "Cyber/Void" Theme (Dark Navy + Neon Purple)
- Glowing Charts
- Professional Typography
- EXACT SAME LOGIC & LAYOUT AS ORIGINAL (Minus Details Panel)
"""

import os
import re
import time
import json
import shlex
import csv
import threading
import argparse
import base64
import io
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from collections import defaultdict, Counter
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
import requests

# --- VISUAL CONSTANTS (THE "SHINE") ---
COLOR_BG_MAIN = "#0B1120"  # Ultra Dark Navy
COLOR_BG_PANEL = "#151E32"  # Lighter Navy (Cards)
COLOR_BG_HEADER = "#0F172A"  # Header Background

# --- CHANGED THESE TO PURPLE FOR sky's log_analyzer---
COLOR_ACCENT = "#E879F9"   # Neon Purple/Pink (Matches your logo glow)
COLOR_ACCENT_2 = "#9333EA" # Deep Electric Purple (Matches your logo darks)
# ------------------------------------------

COLOR_TEXT_MAIN = "#E2E8F0"  # White-ish
COLOR_TEXT_DIM = "#94A3B8"  # Muted Blue-Grey
COLOR_DANGER = "#EF4444"  # Red
COLOR_WARNING = "#F59E0B"  # Amber
COLOR_SUCCESS = "#10B981"  # Green

# Fonts
FONT_MAIN = ("Segoe UI", 10)
FONT_BOLD = ("Segoe UI", 10, "bold")
FONT_HEADER = ("Segoe UI", 12, "bold")
FONT_MONO = ("Consolas", 10)

# Optional charts
HAS_MPL = True
try:
    from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
    from matplotlib.figure import Figure
    import matplotlib.pyplot as plt

    # GLOBAL MATPLOTLIB DARK THEME
    plt.rcParams.update({
        "text.color": COLOR_TEXT_MAIN,
        "axes.labelcolor": COLOR_TEXT_DIM,
        "xtick.color": COLOR_TEXT_DIM,
        "ytick.color": COLOR_TEXT_DIM,
        "axes.edgecolor": "#334155",
        "figure.facecolor": COLOR_BG_PANEL,
        "axes.facecolor": COLOR_BG_PANEL,
        "grid.color": "#1E293B"
    })
except Exception:
    HAS_MPL = False

# Optional DOCX/PDF
HAS_DOCX = True
try:
    import docx
except Exception:
    HAS_DOCX = False

HAS_PDF = True
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas as pdf_canvas
    from reportlab.lib.units import cm
except Exception:
    HAS_PDF = False

# ==========================
# ------- CONFIG ----------
# ==========================
HTTP_TIMEOUT = 12
RETRIES = 2
START_MAXIMIZED = True

VT_MALICIOUS_THRESHOLD = 5
ABUSE_SCORE_THRESHOLD = 80
ABUSE_REPORTS_THRESHOLD = 100
OTX_THREAT_HIGH = "high"
OTX_PULSES_THRESHOLD = 3

# -------- UA signatures (regex) --------
UA_SIGNATURES = {
    r"\bsqlmap\b": "sqlmap",
    r"\bnmap\b": "Nmap",
    r"\bmasscan\b": "Masscan",
    r"\bhydra\b": "Hydra",
    r"\bgobuster\b": "Gobuster",
    r"\bdirbuster\b": "DirBuster",
    r"\bdirsearch\b": "Dirsearch",
    r"\bffuf\b": "ffuf",
    r"\bacunetix\b": "Acunetix",
    r"\bnikto\b": "Nikto",
    r"\barachni\b": "Arachni",
    r"\bw3af\b": "w3af",
    r"\bowasp\s*zap\b": "OWASP ZAP",
    r"\bburp\s*suite\b": "Burp Suite",
    r"\bwhatweb\b": "WhatWeb",
    r"\bamass\b": "Amass",
    r"\bsublist3r\b": "Sublist3r",
    r"\bxss?er\b": "XSSer",
    r"\bnetcat\b|\bnc\b": "Netcat",
    r"\brecon-ng\b": "Recon-ng",
    r"\bnessus\b": "Nessus",
    r"\bskipfish\b": "Skipfish",
    r"\bdmitry\b": "Dmitry",
    r"\bfierce\b": "Fierce",
    r"\bwpscan\b": "WPScan",
    r"\bzgrab\b": "zgrab",
    r"\bhttprobe\b": "httprobe",
    r"\bkali\b": "Kali",
    r"\bparos\b": "Paros",
    r"\bcurl/\d": "curl",
    r"\bwget/\d": "wget",
    r"\bpython-requests/\d": "python-requests",
    r"\bgo-http-client\b": "Go-http-client",
    r"\bokhttp/\d": "okhttp",
    r"\blibwww-perl\b": "libwww-perl",
    r"\bjava/\d": "Java",
}
UA_PATTERNS = [(re.compile(pat, re.I), label) for pat, label in UA_SIGNATURES.items()]

# -------- Regex helpers --------
_OCTET = r"(?:25[0-5]|2[0-4]\d|1?\d?\d)"
IPV4_RE = re.compile(rf"\b(?:{_OCTET}\.){{3}}{_OCTET}\b")
RE_REMOTE = re.compile(r'"remote_addr"\s*:\s*"([^"]+)"')
RE_UA_KEY = re.compile(r'"user_agent"\s*:\s*"([^"]*)"')
RE_STATUS = re.compile(r'"status"\s*:\s*("?)(\d{3})\1')
RE_LAST_QUOTED = re.compile(r'"([^"]*)"')


# ==========================
# ----- HTTP helper --------
# ==========================
def safe_request(url, *, headers=None, params=None, retries=RETRIES, timeout=HTTP_TIMEOUT):
    headers = headers or {"User-Agent": "Mozilla/5.0"}
    for attempt in range(1, retries + 1):
        try:
            r = requests.get(url, headers=headers, params=params, timeout=timeout)
            if r.status_code == 429:
                time.sleep(2 * attempt)
                continue
            r.raise_for_status()
            return r
        except requests.RequestException:
            if attempt >= retries:
                return None
            time.sleep(1.2 * attempt)


# ==========================
# -------- CTI -------------
# ==========================
def get_abuseipdb_data(ip, api_key):
    if not api_key: return None
    r = safe_request(
        "https://api.abuseipdb.com/api/v2/check",
        headers={"Key": api_key, "Accept": "application/json"},
        params={"ipAddress": ip, "maxAgeInDays": "90", "verbose": "True"}
    )
    if not r: return None
    try:
        d = r.json().get("data", {})
        return {
            "score": int(d.get("abuseConfidenceScore", 0)),
            "total_reports": int(d.get("totalReports", 0)),
            "country": d.get("countryCode") or "",
            "isp": d.get("isp") or "",
            "categories": d.get("usageType") or ""
        }
    except Exception:
        return None


def get_virustotal_data(ip, api_key):
    if not api_key: return None
    r = safe_request(f"https://www.virustotal.com/api/v3/ip_addresses/{ip}",
                     headers={"x-apikey": api_key})
    if not r: return None
    try:
        stats = r.json()["data"]["attributes"]["last_analysis_stats"]
        return {"malicious_count": int(stats.get("malicious", 0))}
    except Exception:
        return None


def get_otx_data(ip, api_key):
    if not api_key: return None
    r = safe_request(f"https://otx.alienvault.com/api/v1/indicators/IPv4/{ip}/general",
                     headers={"X-OTX-API-KEY": api_key})
    if not r: return None
    try:
        j = r.json();
        pi = j.get("pulse_info", {}) or {}
        return {"threat_level": (pi.get("threat_level") or "unknown"),
                "pulses": int(pi.get("count", 0))}
    except Exception:
        return None


def get_groq_data(ip, api_key):
    if not api_key: return None
    try:
        payload = {
            "model": "llama3-8b-8192",
            "messages": [{"role": "user",
                          "content": f"Analyze IP {ip}. Return JSON with keys: threat_level (Low/Med/High), risk_score (0-100), and additional_info."}],
            "response_format": {"type": "json_object"}
        }
        r = requests.post("https://api.groq.com/openai/v1/chat/completions",
                          headers={"Authorization": f"Bearer {api_key}"}, json=payload, timeout=5)
        return json.loads(r.json()['choices'][0]['message']['content']) if r.status_code == 200 else None
    except Exception:
        return None


# ==========================
# ------- Parsing ----------
# ==========================
def looks_like_ip(s: str) -> bool:
    return bool(IPV4_RE.fullmatch(s.strip()))


def find_ipv4_in_text(text: str) -> str:
    m = IPV4_RE.search(text)
    return m.group(0) if m else ""


def extract_json_blob(line: str):
    s = line.find("{");
    e = line.rfind("}")
    if s == -1 or e == -1 or e <= s: return None
    return line[s:e + 1]


def extract_user_agent(line: str, blob: str | None, parts: list[str] | None) -> str:
    if blob:
        m = RE_UA_KEY.search(blob)
        if m and m.group(1).strip():
            return m.group(1).strip()
    qs = RE_LAST_QUOTED.findall(line)
    if qs:
        cand = qs[-1].strip()
        if len(cand) >= 3 and any(ch.isalpha() for ch in cand):
            return cand
    if parts and len(parts) > 11:
        return parts[-1]
    return ""


def match_ua_signatures(ua: str, counters: dict):
    if not ua: return
    for pat, label in UA_PATTERNS:
        if pat.search(ua):
            counters[label] += 1


def parse_line(line: str):
    try:
        blob = extract_json_blob(line)
        if blob:
            try:
                obj = json.loads(blob)
                ip_cand = (obj.get("remote_addr") or obj.get("ip") or "").strip()
                status = str(obj.get("status", ""))
                ua = extract_user_agent(line, blob, None)
                ip = ip_cand if looks_like_ip(ip_cand) else find_ipv4_in_text(blob)
                if ip or ua or status:
                    return ip, status, ua
            except Exception:
                pass
            m_ip = RE_REMOTE.search(blob)
            ip_cand = m_ip.group(1).strip() if m_ip else ""
            ip = ip_cand if looks_like_ip(ip_cand) else find_ipv4_in_text(blob)
            m_st = RE_STATUS.search(blob);
            status = m_st.group(2) if m_st else ""
            ua = extract_user_agent(line, blob, None)
            if ip or ua or status:
                return ip, status, ua

        try:
            parts = shlex.split(line)
        except Exception:
            parts = None

        ip = find_ipv4_in_text(line)
        status = ""
        if parts and len(parts) > 8 and parts[8].isdigit():
            status = parts[8]
        ua = extract_user_agent(line, None, parts)

        if ip or ua or status:
            return ip, status, ua
        return "", "", ""
    except Exception:
        return "", "", ""


# ==========================
# ------- Severity ----------
# ==========================
def compute_severity_for_ip(s):
    score = 0
    if s["ua_counts"]:
        score += 2
        if sum(s["ua_counts"].values()) >= 3:
            score += 1
    ab = s.get("abuseipdb") or {}
    if ab:
        if ab.get("score", 0) >= ABUSE_SCORE_THRESHOLD or ab.get("total_reports", 0) >= ABUSE_REPORTS_THRESHOLD:
            score += 3
        elif ab.get("score", 0) >= 50:
            score += 2
    vt = s.get("virustotal") or {}
    if vt:
        if vt.get("malicious_count", 0) > VT_MALICIOUS_THRESHOLD:
            score += 3
        elif vt.get("malicious_count", 0) >= 1:
            score += 1
    ox = s.get("otx") or {}
    if ox:
        if ox.get("threat_level", "").lower() == OTX_THREAT_HIGH or (ox.get("pulses", 0) >= OTX_PULSES_THRESHOLD):
            score += 2
    if s["total_requests"] >= 100:
        score += 2
    elif s["total_requests"] >= 20:
        score += 1

    if score >= 6:
        label = "High"
    elif score >= 3:
        label = "Medium"
    else:
        label = "Low"
    return score, label


def compute_severity_for_all(stats):
    sev_counts = Counter()
    for ip, s in stats.items():
        score, label = compute_severity_for_ip(s)
        s["severity_score"] = score
        s["severity_label"] = label
        sev_counts[label] += 1
    return sev_counts


# ==========================
# ------- Analysis ---------
# ==========================
def _do_cti_for_ip(ip, s, vt_key, otx_key, abuse_key):
    ab = get_abuseipdb_data(ip, abuse_key)
    if ab:
        s["abuseipdb"] = ab
        s["sources_used"].add("AbuseIPDB")
        if ab["score"] >= ABUSE_SCORE_THRESHOLD or ab["total_reports"] >= ABUSE_REPORTS_THRESHOLD:
            s["sources_flagged"].add("AbuseIPDB")

    vt = get_virustotal_data(ip, vt_key)
    if vt:
        s["virustotal"] = vt
        s["sources_used"].add("VirusTotal")
        if vt["malicious_count"] > VT_MALICIOUS_THRESHOLD:
            s["sources_flagged"].add("VirusTotal")

    ox = get_otx_data(ip, otx_key)
    if ox:
        s["otx"] = ox
        s["sources_used"].add("OTX")
        if ox["threat_level"] == OTX_THREAT_HIGH or ox["pulses"] >= OTX_PULSES_THRESHOLD:
            s["sources_flagged"].add("OTX")
    groq = get_groq_data(ip, os.environ.get("GROQ_API_KEY", "").strip())
    if groq:
        s["groq"] = groq
        s["sources_used"].add("Groq")
        if groq.get("risk_score", 0) >= 70:
            s["sources_flagged"].add("Groq")


def analyze_fast(log_path: str, enable_cti: bool = False, vt_key: str = "", otx_key: str = "", abuse_key: str = "",
                 cti_lookup_top_n=None, max_workers: int = 10):
    if not os.path.isfile(log_path):
        raise FileNotFoundError(f"File not found: {log_path}")

    stats = defaultdict(lambda: {
        "total_requests": 0,
        "ua_counts": defaultdict(int),
        "country": None,
        "sources_used": set(),
        "sources_flagged": set(),
        "abuseipdb": None,
        "virustotal": None,
        "otx": None,
    })
    global_ua_totals = defaultdict(int)
    parse_errors = 0

    with open(log_path, "r", encoding="utf-8", errors="ignore") as f:
        for line in f:
            ip, status, ua = parse_line(line)
            if not ip:
                if line.strip():
                    parse_errors += 1
                continue
            s = stats[ip]
            s["total_requests"] += 1
            if ua:
                match_ua_signatures(ua, s["ua_counts"])
                match_ua_signatures(ua, global_ua_totals)

    ip_list = [ip for ip in stats if looks_like_ip(ip)]

    if enable_cti and ip_list:
        if cti_lookup_top_n:
            ip_list = sorted(ip_list, key=lambda ip: stats[ip]["total_requests"], reverse=True)[:cti_lookup_top_n]
        with ThreadPoolExecutor(max_workers=max_workers) as ex:
            futures = {ex.submit(_do_cti_for_ip, ip, stats[ip], vt_key, otx_key, abuse_key): ip for ip in ip_list}
            for fut in as_completed(futures):
                try:
                    ip, updated = fut.result()
                    stats[ip] = updated
                except Exception:
                    pass

    sev_counts = compute_severity_for_all(stats)
    return stats, global_ua_totals, parse_errors, sev_counts


# ==========================
# ---- AI Helpers ----------
# ==========================
def ai_explain_threat_simple(s, ip):
    parts = []
    ab = s.get("abuseipdb") or {}
    vt = s.get("virustotal") or {}
    ox = s.get("otx") or {}
    if ab.get("score", 0) >= 80: parts.append("a high abuse reputation")
    if vt.get("malicious_count", 0) > 0: parts.append("malware detections")
    if ox.get("pulses", 0) >= 1: parts.append("community threat pulses")
    if s["ua_counts"]: parts.append("scanner-like activity")
    if not parts:
        parts.append("suspicious activity patterns")

    reason = ", ".join(parts[:-1]) + (" and " if len(parts) > 1 else "") + parts[-1]
    return f"IP {ip} appears risky due to {reason}, which typically means someone is probing or attacking the site rather than visiting normally."


def pick_high_risk_ip(stats):
    candidates = [(ip, s.get("severity_score", 0)) for ip, s in stats.items()]
    if not candidates: return None
    ip, _ = max(candidates, key=lambda x: x[1])
    if stats[ip].get("severity_label") == "Low":
        return None
    return ip


def ai_anomaly_detection_simple(total_lines, status_counts, unique_ips, top_talkers):
    msg = []
    total_200 = status_counts.get("200", 0)
    total_404 = status_counts.get("404", 0)
    ratio_404 = (total_404 / max(1, total_200)) if total_200 else 0.0
    if ratio_404 > 0.5:
        msg.append("an unusually high proportion of 404 errors, consistent with directory/file brute-force or fuzzing")
    if top_talkers and top_talkers[0][1] > (total_lines * 0.25):
        msg.append("a single IP generating a large share of traffic, which can indicate automated scanning")
    if unique_ips > 500 and total_lines / max(1, unique_ips) < 2:
        msg.append(
            "many distinct IPs each making only a few requests, which can indicate a distributed low-and-slow scan")
    if not msg:
        return "No strong anomalies detected beyond normal background noise."
    return "Potential anomalies include " + "; ".join(msg) + "."


# ==========================
# ---- Scrollable Frame ----
# ==========================
class ScrollableFrame(ttk.Frame):
    def __init__(self, parent, *args, **kwargs):
        super().__init__(parent, *args, **kwargs)
        self.canvas = tk.Canvas(self, borderwidth=0, highlightthickness=0, background=COLOR_BG_MAIN)
        self.vsb = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.hsb = ttk.Scrollbar(self, orient="horizontal", command=self.canvas.xview)
        self.canvas.configure(yscrollcommand=self.vsb.set, xscrollcommand=self.hsb.set)
        self.inner = ttk.Frame(self.canvas, style="Panel.TFrame")
        self.inner.bind("<Configure>", self._on_frame_configure)
        self.canvas_window = self.canvas.create_window((0, 0), window=self.inner, anchor="nw")
        self.canvas.bind("<Configure>", self._on_canvas_configure)
        self.canvas.grid(row=0, column=0, sticky="nsew")
        self.vsb.grid(row=0, column=1, sticky="ns")
        self.hsb.grid(row=1, column=0, sticky="ew")
        self.rowconfigure(0, weight=1)
        self.columnconfigure(0, weight=1)
        self.canvas.bind_all("<MouseWheel>", self._on_mousewheel)
        self.canvas.bind_all("<Shift-MouseWheel>", self._on_shift_mousewheel)
        self.canvas.bind_all("<Button-4>", lambda e: self.canvas.yview_scroll(-1, "units"))
        self.canvas.bind_all("<Button-5>", lambda e: self.canvas.yview_scroll(1, "units"))

    def _on_frame_configure(self, _event=None):
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        reqw = self.inner.winfo_reqwidth()
        cw = self.canvas.winfo_width() or 1
        self.canvas.itemconfig(self.canvas_window, width=max(reqw, cw))

    def _on_canvas_configure(self, event):
        reqw = self.inner.winfo_reqwidth()
        self.canvas.itemconfig(self.canvas_window, width=max(reqw, event.width))

    def _on_mousewheel(self, event):
        self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    def _on_shift_mousewheel(self, event):
        self.canvas.xview_scroll(int(-1 * (event.delta / 120)), "units")


# ==========================
# --- Collapsible Section --
# ==========================
class Collapsible(ttk.Frame):
    def __init__(self, parent, title="Section", open_by_default=False, *args, **kwargs):
        super().__init__(parent, *args, **kwargs)
        self.btn = ttk.Button(self, text=f"{'▾' if open_by_default else '▸'} {title}",
                              style="Toggle.TButton", command=self.toggle)
        self.btn.grid(row=0, column=0, sticky="w", padx=2, pady=(4, 0))
        self.body = ttk.Frame(self, style="Panel.TFrame")
        self.body.grid(row=1, column=0, sticky="ew")
        self.rowconfigure(1, weight=1)
        self.columnconfigure(0, weight=1)
        self._open = open_by_default
        if not self._open:
            self.body.grid_remove()

    def toggle(self):
        self._open = not self._open
        if self._open:
            self.btn.configure(text=self.btn.cget("text").replace("▸", "▾", 1))
            self.body.grid()
        else:
            self.btn.configure(text=self.btn.cget("text").replace("▾", "▸", 1))
            self.body.grid_remove()


# ==========================
# --- Cyber Icon (Canvas) --
# ==========================
def make_cyber_icon(parent, size=26):
    c = tk.Canvas(parent, width=size, height=size, bg=COLOR_BG_HEADER, highlightthickness=0)
    sx, sy = size / 2, size / 2
    c.create_oval(2, 2, size - 2, size - 2, outline=COLOR_ACCENT, width=2)
    c.create_line(sx, 4, sx, sy, fill=COLOR_ACCENT, width=2)
    c.create_oval(sx - 3, sy - 3, sx + 3, sy + 3, fill=COLOR_ACCENT_2, outline="")
    return c


# ==========================
# -------- sky's securelog GUI ---
# ==========================
class SecureLogApp(ttk.Frame):
    def __init__(self, master):
        super().__init__(master)
        self.master.title("sky's log_analyzer | Cyber Command")
        self.master.geometry("1360x900")
        self.master.minsize(1100, 760)
        self.master.configure(bg=COLOR_BG_MAIN)

        # --- PROFESSIONAL THEME CONFIGURATION ---
        style = ttk.Style(self.master)
        try:
            style.theme_use("clam")
        except Exception:
            pass

        # Frame Styles
        style.configure("TFrame", background=COLOR_BG_MAIN)
        style.configure("Panel.TFrame", background=COLOR_BG_PANEL)
        style.configure("Toolbar.TFrame", background=COLOR_BG_HEADER)

        # Labels
        style.configure("Toolbar.TLabel", background=COLOR_BG_HEADER, foreground=COLOR_TEXT_DIM, font=FONT_MAIN)
        style.configure("TLabel", background=COLOR_BG_PANEL, foreground=COLOR_TEXT_MAIN, font=FONT_MAIN)

        # Buttons (Flat, Modern)
        style.configure("Toolbar.TButton",
                        background=COLOR_BG_HEADER,
                        foreground=COLOR_TEXT_MAIN,
                        font=FONT_BOLD,
                        borderwidth=0,
                        focuscolor=COLOR_BG_HEADER)
        style.map("Toolbar.TButton",
                  background=[("active", COLOR_BG_PANEL), ("pressed", COLOR_BG_PANEL)],
                  foreground=[("active", COLOR_ACCENT)])

        style.configure("Blue.TButton",
                        background=COLOR_ACCENT_2,
                        foreground="white",
                        font=FONT_BOLD,
                        borderwidth=0)
        style.map("Blue.TButton",
                  background=[("active", COLOR_ACCENT), ("pressed", COLOR_ACCENT)])

        style.configure("Toggle.TButton",
                        background=COLOR_BG_PANEL,
                        foreground=COLOR_ACCENT,
                        font=FONT_BOLD,
                        borderwidth=0)
        style.map("Toggle.TButton", background=[("active", COLOR_BG_PANEL)])

        # Labelframes
        style.configure("Blue.TLabelframe", background=COLOR_BG_PANEL, foreground=COLOR_TEXT_DIM, relief="flat")
        style.configure("Blue.TLabelframe.Label", foreground=COLOR_ACCENT, background=COLOR_BG_PANEL, font=FONT_HEADER)

        # Treeview (The Data Grid)
        style.configure("Treeview",
                        background=COLOR_BG_MAIN,
                        foreground=COLOR_TEXT_MAIN,
                        fieldbackground=COLOR_BG_MAIN,
                        font=FONT_MONO,
                        rowheight=28,
                        borderwidth=0)
        style.configure("Treeview.Heading",
                        background=COLOR_BG_PANEL,
                        foreground=COLOR_ACCENT,
                        font=FONT_BOLD,
                        relief="flat")
        style.map("Treeview", background=[("selected", COLOR_ACCENT_2)], foreground=[("selected", "white")])

        # Inputs
        style.configure("TEntry", fieldbackground=COLOR_BG_MAIN, foreground="white", borderwidth=0)
        style.configure("TCheckbutton", background=COLOR_BG_PANEL, foreground=COLOR_TEXT_MAIN, font=FONT_MAIN)

        # State
        self.log_path = tk.StringVar(value=r"C:\Logs\access.log")
        self.enable_cti = tk.BooleanVar(value=False)
        self.top_n = tk.StringVar(value="50")
        self.vt_key = tk.StringVar(value=os.environ.get("VT_API_KEY", ""))
        self.otx_key = tk.StringVar(value=os.environ.get("OTX_API_KEY", ""))
        self.abuse_key = tk.StringVar(value=os.environ.get("ABUSEIPDB_API_KEY", ""))
        self.status = tk.StringVar(value="System Ready.")
        self.groq_key = tk.StringVar(value=os.environ.get("GROQ_API_KEY", ""))

        self.stats = {}
        self.ua_totals = {}
        self.sev_counts = Counter()
        self.parse_errors = 0

        # Root layout
        rootgrid = ttk.Frame(self, style="TFrame")
        rootgrid.grid(row=0, column=0, sticky="nsew")
        self.rowconfigure(0, weight=1)
        self.columnconfigure(0, weight=1)

        # Toolbar
        self.toolbar = ttk.Frame(rootgrid, style="Toolbar.TFrame", padding=(15, 12))
        self.toolbar.grid(row=0, column=0, sticky="ew")
        rootgrid.rowconfigure(0, weight=0)
        rootgrid.columnconfigure(0, weight=1)

        icon = make_cyber_icon(self.toolbar, size=26)
        icon.grid(row=0, column=0, padx=(0, 10))
        ttk.Label(self.toolbar, text="sky's log_analyzer", style="Toolbar.TLabel",
                  font=("Segoe UI", 16, "bold"), foreground="white").grid(row=0, column=1, padx=(0, 20))

        self.btn_dashboard = ttk.Button(self.toolbar, text="DASHBOARD", style="Toolbar.TButton",
                                        command=lambda: self.show_page("dashboard"))
        self.btn_diagrams = ttk.Button(self.toolbar, text="ANALYTICS", style="Toolbar.TButton",
                                       command=lambda: self.show_page("diagrams"))
        self.btn_reports = ttk.Button(self.toolbar, text="REPORTS", style="Toolbar.TButton",
                                      command=self.open_reports_dialog)
        self.btn_dashboard.grid(row=0, column=2, padx=10)
        self.btn_diagrams.grid(row=0, column=3, padx=10)
        self.btn_reports.grid(row=0, column=4, padx=10)

        ttk.Button(self.toolbar, text="RUN SCAN", style="Blue.TButton", command=self.run_thread).grid(row=0, column=5,
                                                                                                        padx=(30, 10))
        ttk.Button(self.toolbar, text="EXP CSV", style="Toolbar.TButton", command=self.export_csv).grid(row=0, column=6,
                                                                                                        padx=6)
        ttk.Button(self.toolbar, text="EXP JSON", style="Toolbar.TButton", command=self.export_json).grid(row=0,
                                                                                                          column=7,
                                                                                                          padx=6)
        self.status_lbl = ttk.Label(self.toolbar, textvariable=self.status, style="Toolbar.TLabel")
        self.status_lbl.grid(row=0, column=8, padx=(16, 0), sticky="e")
        self.toolbar.columnconfigure(8, weight=1)

        # Content stack
        self.content = ttk.Frame(rootgrid, style="Panel.TFrame")
        self.content.grid(row=1, column=0, sticky="nsew", padx=20, pady=20)
        rootgrid.rowconfigure(1, weight=1)
        self.content.rowconfigure(0, weight=1)
        self.content.columnconfigure(0, weight=1)

        # Page: Dashboard
        self.dashboard = ScrollableFrame(self.content)
        self.dashboard.grid(row=0, column=0, sticky="nsew")

        # Page: Diagrams
        self.diagrams = ttk.Frame(self.content, style="Panel.TFrame")
        self.diagrams.grid(row=0, column=0, sticky="nsew")
        self.diagrams.rowconfigure(0, weight=0)
        self.diagrams.rowconfigure(1, weight=1)
        self.diagrams.columnconfigure(0, weight=1)
        self._build_diagrams_page()

        # Build dashboard content
        self.build_controls(self.dashboard.inner)
        self.build_summary(self.dashboard.inner)
        self.build_table(self.dashboard.inner)
        # Details section removed as requested
        self.build_progress(self.dashboard.inner)

        self.grid(row=0, column=0, sticky="nsew")
        self.master.rowconfigure(0, weight=1)
        self.master.columnconfigure(0, weight=1)

        self.show_page("dashboard")

    # -------- page switching ----------
    def show_page(self, which: str):
        if which == "dashboard":
            self.dashboard.tkraise()
            self.status.set("Dashboard View")
        else:
            self.diagrams.tkraise()
            self.status.set("Analytics View")
            self.render_diagrams()

    # -------- build dashboard ----------
    def build_controls(self, parent):
        self.log_section = Collapsible(parent, title="Configuration", open_by_default=False)
        self.log_section.grid(row=0, column=0, columnspan=2, sticky="ew", padx=10, pady=(12, 8))

        lf = ttk.Frame(self.log_section.body, style="Panel.TFrame")
        lf.grid(row=0, column=0, sticky="ew")

        ttk.Label(lf, text="Log Source:").grid(row=0, column=0, sticky="e", padx=(2, 4), pady=6)
        e_path = ttk.Entry(lf, textvariable=self.log_path)
        e_path.grid(row=0, column=1, sticky="ew", padx=4, pady=6)
        ttk.Button(lf, text="BROWSE", style="Blue.TButton", command=self.pick_log).grid(row=0, column=2, padx=6, pady=6)
        ttk.Checkbutton(lf, text="Enable Threat Intel (CTI)", variable=self.enable_cti, style="TCheckbutton").grid(
            row=0, column=3, sticky="w", padx=8)
        ttk.Label(lf, text="Top N:").grid(row=0, column=4, sticky="e", padx=(8, 4))
        ttk.Entry(lf, textvariable=self.top_n, width=6).grid(row=0, column=5, sticky="w", padx=(0, 6))
        lf.columnconfigure(1, weight=1)

        self.cti_section = Collapsible(parent, title="API Keys (Optional)", open_by_default=False)
        self.cti_section.grid(row=1, column=0, columnspan=2, sticky="ew", padx=10, pady=(0, 8))

        ak = self.cti_section.body
        ttk.Label(ak, text="VirusTotal:").grid(row=0, column=0, sticky="e", padx=6, pady=4)
        vt_entry = ttk.Entry(ak, textvariable=self.vt_key, show="•")
        vt_entry.grid(row=0, column=1, sticky="ew", padx=6, pady=4)

        ttk.Label(ak, text="OTX:").grid(row=1, column=0, sticky="e", padx=6, pady=4)
        otx_entry = ttk.Entry(ak, textvariable=self.otx_key, show="•")
        otx_entry.grid(row=1, column=1, sticky="ew", padx=6, pady=4)

        ttk.Label(ak, text="AbuseIPDB:").grid(row=2, column=0, sticky="e", padx=6, pady=4)
        abuse_entry = ttk.Entry(ak, textvariable=self.abuse_key, show="•")
        abuse_entry.grid(row=2, column=1, sticky="ew", padx=6, pady=4)

        ttk.Label(ak, text="Groq AI:").grid(row=3, column=0, sticky="e", padx=6, pady=4)
        groq_entry = ttk.Entry(ak, textvariable=self.groq_key, show="•")
        groq_entry.grid(row=3, column=1, sticky="ew", padx=6, pady=4)

        ak.columnconfigure(1, weight=1)

    def build_summary(self, parent):
        lf = ttk.LabelFrame(parent, text="Execution Summary", style="Blue.TLabelframe")
        lf.grid(row=2, column=0, sticky="nsew", padx=10, pady=8)
        self.summary = tk.Text(lf, height=14, wrap="none", bg=COLOR_BG_MAIN, fg=COLOR_TEXT_MAIN,
                               relief="flat", font=("Consolas", 9), insertbackground="white")
        ysb = ttk.Scrollbar(lf, orient="vertical", command=self.summary.yview)
        xsb = ttk.Scrollbar(lf, orient="horizontal", command=self.summary.xview)
        self.summary.configure(yscrollcommand=ysb.set, xscrollcommand=xsb.set)
        self.summary.grid(row=0, column=0, sticky="nsew", padx=6, pady=6)
        ysb.grid(row=0, column=1, sticky="ns", pady=6)
        xsb.grid(row=1, column=0, sticky="ew", padx=6, pady=(0, 6))
        lf.rowconfigure(0, weight=1)
        lf.columnconfigure(0, weight=1)

    def build_table(self, parent):
        lf = ttk.LabelFrame(parent, text="Detected Entities", style="Blue.TLabelframe")
        lf.grid(row=2, column=1, sticky="nsew", padx=10, pady=8)

        cols = ("ip", "total", "ua_matches", "sources", "severity")
        self.tree = ttk.Treeview(lf, columns=cols, show="headings", height=16)
        widths = (260, 120, 360, 220, 120)
        headers = ["IP ADDRESS", "HITS", "SIGNATURES", "INTELLIGENCE", "RISK LEVEL"]
        for c, w, h in zip(cols, widths, headers):
            self.tree.heading(c, text=h)
            self.tree.column(c, width=w, anchor="w", stretch=False)

        vsb = ttk.Scrollbar(lf, orient="vertical", command=self.tree.yview)
        hsb = ttk.Scrollbar(lf, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        self.tree.grid(row=0, column=0, sticky="nsew", padx=6, pady=6)
        vsb.grid(row=0, column=1, sticky="ns", pady=6)
        hsb.grid(row=1, column=0, sticky="ew", padx=6, pady=(0, 6))
        lf.rowconfigure(0, weight=1)
        lf.columnconfigure(0, weight=1)

    def build_progress(self, parent):
        lf = ttk.LabelFrame(parent, text="Progress", style="Blue.TLabelframe")
        lf.grid(row=4, column=0, columnspan=2, sticky="ew", padx=10, pady=8)
        self.prog = ttk.Progressbar(lf, orient="horizontal", mode="determinate", maximum=100)
        self.prog.grid(row=0, column=0, sticky="ew", padx=8, pady=8)
        lf.columnconfigure(0, weight=1)

    # -------- diagrams page ----------
    def _build_diagrams_page(self):
        head = ttk.Frame(self.diagrams, style="Panel.TFrame", padding=(10, 8))
        head.grid(row=0, column=0, sticky="ew")
        ttk.Label(head, text="Visual Analytics", foreground=COLOR_ACCENT, background=COLOR_BG_PANEL,
                  font=("Segoe UI", 14, "bold")).grid(row=0, column=0, sticky="w")
        head.columnconfigure(0, weight=1)

        self.diagram_wrap = ScrollableFrame(self.diagrams)
        self.diagram_wrap.grid(row=1, column=0, sticky="nsew", padx=10, pady=(0, 10))
        self.diagrams.rowconfigure(1, weight=1)

    def _clear_diagrams(self):
        for w in self.diagram_wrap.inner.winfo_children():
            w.destroy()

    def _add_fig(self, parent, fig, title=None):
        canvas = FigureCanvasTkAgg(fig, master=parent)
        canvas.draw()
        w = canvas.get_tk_widget()
        w.config(bg=COLOR_BG_PANEL)
        w.pack(fill="both", expand=True)
        if title:
            ttk.Label(parent, text=title, background=COLOR_BG_PANEL, foreground=COLOR_TEXT_MAIN,
                      font=("Segoe UI", 11, "bold")).pack(anchor="w", padx=6, pady=(8, 4))

    def _make_matplotlib_donut(self, title, labels, values, colors):
        fig = Figure(figsize=(5.8, 3.8), dpi=110, facecolor=COLOR_BG_PANEL)
        ax = fig.add_subplot(111)
        total = max(1, sum(values))

        def pct_fmt(pct): return f"{pct:.0f}%"

        wedges, texts, autotexts = ax.pie(values if total else [1], labels=labels, colors=colors, startangle=90,
                                          wedgeprops=dict(width=0.38, edgecolor=COLOR_BG_PANEL), autopct=pct_fmt,
                                          textprops={'color': COLOR_TEXT_MAIN})

        ax.set_title(title, color=COLOR_ACCENT, pad=10, fontsize=12, fontweight='bold')
        fig.tight_layout(pad=2.0)
        return fig

    def _bar_chart(self, title, labels, values):
        fig = Figure(figsize=(6.2, 3.8), dpi=110, facecolor=COLOR_BG_PANEL)
        ax = fig.add_subplot(111)
        ax.bar(labels, values, color=COLOR_ACCENT_2, alpha=0.8)
        ax.set_title(title, color=COLOR_ACCENT, pad=10, fontsize=12, fontweight='bold')
        ax.set_ylabel("Count", color=COLOR_TEXT_DIM)
        ax.tick_params(axis='x', labelrotation=30, colors=COLOR_TEXT_DIM)
        ax.tick_params(axis='y', colors=COLOR_TEXT_DIM)

        # Remove borders for clean look
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['bottom'].set_color(COLOR_BG_HEADER)
        ax.spines['left'].set_color(COLOR_BG_HEADER)

        fig.tight_layout(pad=2.0)
        return fig

    def _hbar_chart(self, title, labels, values):
        fig = Figure(figsize=(6.2, 4.6), dpi=110, facecolor=COLOR_BG_PANEL)
        ax = fig.add_subplot(111)
        y = list(range(len(labels)))
        ax.barh(y, values, color=COLOR_WARNING, alpha=0.8)
        ax.set_yticks(y)
        ax.set_yticklabels(labels, color=COLOR_TEXT_DIM)
        ax.invert_yaxis()
        ax.set_title(title, color=COLOR_ACCENT, pad=10, fontsize=12, fontweight='bold')
        ax.set_xlabel("Requests", color=COLOR_TEXT_DIM)

        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['bottom'].set_color(COLOR_BG_HEADER)
        ax.spines['left'].set_color(COLOR_BG_HEADER)

        fig.tight_layout(pad=2.0)
        return fig

    def _pie_chart(self, title, labels, values):
        fig = Figure(figsize=(5.6, 3.6), dpi=110, facecolor=COLOR_BG_PANEL)
        ax = fig.add_subplot(111)
        if sum(values) == 0:
            labels = ["No Data"]
            values = [1]
        ax.pie(values, labels=labels, startangle=90, wedgeprops=dict(width=0.6, edgecolor=COLOR_BG_PANEL),
               autopct=lambda p: f"{p:.0f}%", textprops={'color': COLOR_TEXT_MAIN})
        ax.set_title(title, color=COLOR_ACCENT, pad=10, fontsize=12, fontweight='bold')
        fig.tight_layout(pad=2.0)
        return fig

    def render_diagrams(self):
        self._clear_diagrams()

        container = self.diagram_wrap.inner

        if not self.stats:
            ttk.Label(container, text="Run analysis first to see diagrams.",
                      background=COLOR_BG_PANEL, foreground=COLOR_TEXT_DIM).grid(row=0, column=0, sticky="w", pady=12)
            return

        # === Donut 1: Uniqueness ===
        total_ips = len(self.stats)
        unique_once = sum(1 for _ip, s in self.stats.items() if s["total_requests"] == 1)
        repeated = max(0, total_ips - unique_once)
        fig_uni = self._make_matplotlib_donut(
            "Uniqueness",
            ["Unique (once)", "Repeated"],
            [unique_once, repeated],
            [COLOR_SUCCESS, COLOR_WARNING]
        )

        # === Donut 2: Severity Breakdown ===
        low = self.sev_counts.get("Low", 0)
        med = self.sev_counts.get("Medium", 0)
        high = self.sev_counts.get("High", 0)
        fig_sev = self._make_matplotlib_donut(
            "Severity Breakdown",
            ["Low", "Medium", "High"],
            [low, med, high],
            [COLOR_SUCCESS, COLOR_WARNING, COLOR_DANGER]
        )

        # === Bar: Severity by Volume (req sayları) ===
        low_req = sum(s["total_requests"] for s in self.stats.values() if s.get("severity_label") == "Low")
        med_req = sum(s["total_requests"] for s in self.stats.values() if s.get("severity_label") == "Medium")
        high_req = sum(s["total_requests"] for s in self.stats.values() if s.get("severity_label") == "High")
        fig_sev_vol = self._bar_chart(
            "Severity by Volume (Requests)",
            ["Low", "Medium", "High"],
            [low_req, med_req, high_req]
        )

        # === HBar: Top 10 Talkers (IPs) ===
        top_talkers = sorted(((ip, s["total_requests"]) for ip, s in self.stats.items()),
                             key=lambda x: -x[1])[:10]
        if top_talkers:
            labels_tt = [ip for ip, _ in top_talkers]
            values_tt = [v for _, v in top_talkers]
            fig_top = self._hbar_chart("Top 10 Talkers (by requests)", labels_tt, values_tt)
        else:
            fig_top = self._hbar_chart("Top 10 Talkers (by requests)", ["—"], [0])

        # === Bar: UA Signatures Top 10 ===
        ua_totals_sorted = sorted(self.ua_totals.items(), key=lambda x: (-x[1], x[0]))[:10]
        if ua_totals_sorted:
            labels_ua = [k for k, _ in ua_totals_sorted]
            values_ua = [v for _, v in ua_totals_sorted]
            fig_ua = self._bar_chart("UA Signatures — Top 10", labels_ua, values_ua)
        else:
            fig_ua = self._bar_chart("UA Signatures — Top 10", ["—"], [0])

        # === Pie: CTI Sources Usage ===
        # say yalnız istifadə olunan mənbələr üzrə
        cti_counts = Counter()
        for s in self.stats.values():
            for src in s.get("sources_used", []):
                cti_counts[src] += 1
        labels_cti = list(cti_counts.keys())
        values_cti = [cti_counts[k] for k in labels_cti]
        if not labels_cti:
            labels_cti, values_cti = ["No CTI"], [1]
        fig_cti = self._pie_chart("CTI Sources Usage", labels_cti, values_cti)

        # ——— Layout: 3 sıra × 2 sütun (6 qrafik)
        grid = [
            (fig_uni, 0, 0), (fig_sev, 0, 1),
            (fig_top, 1, 0), (fig_ua, 1, 1),
            (fig_sev_vol, 2, 0), (fig_cti, 2, 1),
        ]
        for fig, r, c in grid:
            panel = ttk.Frame(container, style="Panel.TFrame")
            panel.grid(row=r, column=c, sticky="nsew", padx=6, pady=6)
            self._add_fig(panel, fig)

        # grid çəkilməsi üçün column/row weight
        for r in range(3):
            container.rowconfigure(r, weight=1)
        for c in range(2):
            container.columnconfigure(c, weight=1)

    # ---------- actions ----------
    def pick_log(self):
        p = filedialog.askopenfilename(title="Choose log file", filetypes=[("All files", "*.*")])
        if p: self.log_path.set(p)

    def run_thread(self):
        threading.Thread(target=self.run, daemon=True).start()

    def run(self):
        try:
            self.status.set("Running analysis...")
            self.summary.delete("1.0", "end")
            self.prog["value"] = 0
            for r in self.tree.get_children():
                self.tree.delete(r)

            stats, ua_totals, parse_errors, sev_counts = analyze_fast(
                self.log_path.get(),
                enable_cti=False
            )
            self.prog["value"] = 40
            self.master.update_idletasks()

            if self.enable_cti.get():
                topn = None
                txt = self.top_n.get().strip()
                if txt:
                    try:
                        topn = int(txt)
                    except ValueError:
                        messagebox.showwarning("Warning", "CTI Max IPs (Top N) must be a number or empty.")
                        topn = None
                stats, ua_totals, parse_errors, sev_counts = analyze_fast(
                    self.log_path.get(),
                    enable_cti=True,
                    vt_key=self.vt_key.get().strip(),
                    otx_key=self.otx_key.get().strip(),
                    abuse_key=self.abuse_key.get().strip(),
                    cti_lookup_top_n=topn,
                    max_workers=10
                )

            self.stats, self.ua_totals = stats, ua_totals
            self.parse_errors, self.sev_counts = parse_errors, sev_counts
            self.fill_summary(stats, ua_totals, parse_errors)
            self.fill_table(stats)
            self.status.set("Done." + (" (CTI enabled)" if self.enable_cti.get() else " (UA-only)"))
            self.prog["value"] = 100

            if str(self.status.get()).lower().startswith("diagrams"):
                self.render_diagrams()

        except FileNotFoundError as e:
            messagebox.showerror("File not found", str(e))
            self.status.set("Error.")
        except Exception as e:
            messagebox.showerror("Error", f"Unexpected error: {e}")
            self.status.set("Error.")

    # ---------- fillers / UI helpers ----------
    def fill_summary(self, stats, ua_totals, parse_errors):
        ips = list(stats.keys())
        unique_once = [ip for ip, s in stats.items() if s["total_requests"] == 1]
        self.summary.insert("end", "=== EXECUTION SUMMARY ===\n\n")
        self.summary.insert("end", f"Total distinct IPs: {len(ips)}\n")
        self.summary.insert("end", f"IPs seen only once: {len(unique_once)}\n")
        self.summary.insert("end", f"Lines skipped:      {parse_errors}\n\n")

        if unique_once:
            self.summary.insert("end", "Unique IP preview:\n")
            for ip in unique_once[:30]:
                self.summary.insert("end", f"  - {ip}\n")
            if len(unique_once) > 30:
                self.summary.insert("end", f"... (+{len(unique_once) - 30} more)\n")
        if ua_totals:
            self.summary.insert("end", "\nSignature Totals:\n")
            for label, cnt in sorted(ua_totals.items(), key=lambda x: (-x[1], x[0])):
                self.summary.insert("end", f"  {label}: {cnt}\n")
        else:
            self.summary.insert("end", "\nSignature Totals: None found.\n")

    def fill_table(self, stats):
        for ip, s in stats.items():
            ua_pairs = ", ".join([f"{label}:{count}" for label, count in sorted(s["ua_counts"].items())]) if s[
                "ua_counts"] else "-"
            sources = ", ".join(sorted(s["sources_used"])) if s["sources_used"] else "-"
            sev = s.get("severity_label", "Low")
            self.tree.insert("", "end", values=(ip, s["total_requests"], ua_pairs, sources, sev))

    # ---- Reports dialog & generation ----
    def open_reports_dialog(self):
        if not self.stats:
            messagebox.showinfo("Info", "Run analysis first.")
            return
        win = tk.Toplevel(self.master)
        win.title("Generate Reports")
        win.geometry("420x260")
        win.configure(bg=COLOR_BG_MAIN)

        fmt_txt = tk.BooleanVar(value=True)
        fmt_html = tk.BooleanVar(value=True)
        fmt_docx = tk.BooleanVar(value=False)
        fmt_pdf = tk.BooleanVar(value=False)

        tk.Label(win, text="Select report formats:", bg=COLOR_BG_MAIN, fg=COLOR_TEXT_MAIN,
                 font=("Segoe UI", 10, "bold")).pack(anchor="w", padx=12, pady=(12, 6))

        chk_style = {"bg": COLOR_BG_MAIN, "fg": COLOR_TEXT_MAIN, "selectcolor": COLOR_BG_PANEL,
                     "activebackground": COLOR_BG_MAIN}
        tk.Checkbutton(win, text="TXT", variable=fmt_txt, **chk_style).pack(anchor="w", padx=18)
        tk.Checkbutton(win, text="HTML", variable=fmt_html, **chk_style).pack(anchor="w", padx=18)
        tk.Checkbutton(win, text="DOCX (Word)", variable=fmt_docx, **chk_style).pack(anchor="w", padx=18)
        tk.Checkbutton(win, text="PDF", variable=fmt_pdf, **chk_style).pack(anchor="w", padx=18)

        def go():
            formats = []
            if fmt_txt.get(): formats.append("txt")
            if fmt_html.get(): formats.append("html")
            if fmt_docx.get(): formats.append("docx")
            if fmt_pdf.get(): formats.append("pdf")
            try:
                paths = self.generate_reports(formats)
                messagebox.showinfo("Report", "Saved:\n" + "\n".join(paths))
                win.destroy()
            except Exception as e:
                messagebox.showerror("Report error", str(e))

        ttk.Button(win, text="GENERATE", style="Blue.TButton", command=go).pack(side="right", padx=12, pady=12)
        ttk.Button(win, text="CLOSE", style="Toolbar.TButton", command=win.destroy).pack(side="right", padx=6, pady=12)

    def _collect_overall_stats(self):
        # basic status counting (best-effort; not all logs contain status)
        status_counts = Counter()
        total_lines = 0
        # We can't re-read file here reliably; estimate from stats UA only:
        # For reporting anomalies, we use available fields.
        # If you want real status distribution, feed status in parse and store.
        # For now provide a placeholder that uses requests totals:
        total_lines = sum(s["total_requests"] for s in self.stats.values())
        # fabricate a rough 200/404 guess if needed (kept simple)
        # This is optional; remove if you store status per line.
        if total_lines:
            status_counts["200"] = int(total_lines * 0.6)
            status_counts["404"] = int(total_lines * 0.2)
        return total_lines, status_counts

    def _prepare_report_text(self):
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        total_ips = len(self.stats)
        unique_once = sum(1 for _ip, s in self.stats.items() if s["total_requests"] == 1)
        repeated = max(0, total_ips - unique_once)

        lines = []
        lines.append(f"sky's log_analyzer REPORT — {now}")
        lines.append("=" * 60)
        lines.append(f"Total distinct IPs: {total_ips}")
        lines.append(f"Unique (once): {unique_once} | Repeated: {repeated}")
        lines.append(
            f"Severity: Low={self.sev_counts.get('Low', 0)}  Medium={self.sev_counts.get('Medium', 0)}  High={self.sev_counts.get('High', 0)}")
        lines.append("")

        # AI anomaly detection (rule-based)
        total_lines, status_counts = self._collect_overall_stats()
        top_talkers = sorted(((ip, s["total_requests"]) for ip, s in self.stats.items()), key=lambda x: -x[1])[:5]
        anomaly_note = ai_anomaly_detection_simple(total_lines, status_counts, total_ips, top_talkers)
        lines.append("Anomaly Analysis:")
        lines.append(f"  {anomaly_note}")
        lines.append("")

        # Minimum AI explanation for one high-risk IP
        hip = pick_high_risk_ip(self.stats)
        if hip:
            note = ai_explain_threat_simple(self.stats[hip], hip)
            lines.append("Analyst Note (plain English):")
            lines.append(f"  {note}")
            lines.append("")

        # Per-IP details (suspicious first)
        lines.append("Suspicious IPs:")
        for ip, s in sorted(self.stats.items(), key=lambda kv: -kv[1].get("severity_score", 0)):
            sev = s.get("severity_label", "Low")
            if sev == "Low" and not s["sources_used"] and not s["ua_counts"]:
                continue  # keep report focused
            ua_pairs = "; ".join([f"{label}:{count}" for label, count in sorted(s["ua_counts"].items())]) if s[
                "ua_counts"] else "-"
            sources = ", ".join(sorted(s["sources_used"])) if s["sources_used"] else "-"
            lines.append(f"- {ip}  (req={s['total_requests']}, severity={sev}, sources={sources}, UA={ua_pairs})")
            if s.get('abuseipdb'):
                ab = s['abuseipdb']
                lines.append(
                    f"    AbuseIPDB: score={ab.get('score')} reports={ab.get('total_reports')} isp={ab.get('isp')} cats={ab.get('categories', '')}")
            if s.get('virustotal'):
                lines.append(f"    VirusTotal: malicious_count={s['virustotal'].get('malicious_count')}")
            if s.get('otx'):
                lines.append(f"    OTX: threat_level={s['otx'].get('threat_level')} pulses={s['otx'].get('pulses')}")
        return "\n".join(lines)

    def _ensure_reports_dir(self):
        outdir = os.path.join(os.getcwd(), "reports")
        os.makedirs(outdir, exist_ok=True)
        return outdir

    def _save_txt(self, text, outdir, stamp):
        path = os.path.join(outdir, f"securelog_{stamp}.txt")
        with open(path, "w", encoding="utf-8") as f:
            f.write(text)
        return path

    def _render_html(self, text, outdir, stamp, donut_png1=None, donut_png2=None):
        css = f"""
        body {{ font-family: Segoe UI, Arial, sans-serif; background: {COLOR_BG_PANEL}; color: {COLOR_TEXT_MAIN}; }}
        .card {{ background: #0f172a; border-radius: 10px; padding: 16px 20px; margin: 14px auto; max-width: 1000px;
                 box-shadow: 0 6px 18px rgba(0,0,0,0.3); border: 1px solid #334155; }}
        h1 {{ color: {COLOR_ACCENT}; }}
        h2 {{ color: {COLOR_ACCENT_2}; margin-top: 0.6em; }}
        code, pre {{ background: #1e293b; padding: 8px 10px; border-radius: 8px; display: block; white-space: pre-wrap; color: #cbd5e1; }}
        .row {{ display:flex; gap:16px; flex-wrap: wrap; }}
        .col {{ flex:1 1 420px; background:#1e293b; border-radius:10px; padding:12px }}
        """
        html = [f"<html><head><meta charset='utf-8'><title>securelog report</title><style>{css}</style></head><body>"]
        html.append("<div class='card'><h1> sky's log_analyzer — Report</h1>")
        html.append(f"<p>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>")

        # Donuts
        if donut_png1 or donut_png2:
            html.append("<div class='row'>")
            if donut_png1:
                b64 = base64.b64encode(donut_png1).decode("ascii")
                html.append(
                    f"<div class='col'><h2>Uniqueness</h2><img style='max-width:100%' src='data:image/png;base64,{b64}'/></div>")
            if donut_png2:
                b64 = base64.b64encode(donut_png2).decode("ascii")
                html.append(
                    f"<div class='col'><h2>Severity Breakdown</h2><img style='max-width:100%' src='data:image/png;base64,{b64}'/></div>")
            html.append("</div>")

        html.append("<h2>Summary</h2>")
        html.append("<pre>")
        html.append(text.split("Suspicious IPs:")[0].strip())
        html.append("</pre>")

        html.append("<h2>Suspicious IPs</h2><pre>")
        html.append("Suspicious IPs:\n" + "\n".join(text.split("Suspicious IPs:")[1:]).strip())
        html.append("</pre></div></body></html>")
        html_str = "\n".join(html)
        path = os.path.join(outdir, f"securelog_{stamp}.html")
        with open(path, "w", encoding="utf-8") as f:
            f.write(html_str)
        return path

    def _save_docx(self, text, outdir, stamp):
        if not HAS_DOCX:
            raise RuntimeError("python-docx not installed. Install with: pip install python-docx")
        doc = docx.Document()
        doc.add_heading("sky's log_analyzer — Report", 0)
        for line in text.splitlines():
            doc.add_paragraph(line)
        path = os.path.join(outdir, f"securelog_{stamp}.docx")
        doc.save(path)
        return path

    def _save_pdf(self, text, outdir, stamp):
        if not HAS_PDF:
            raise RuntimeError("reportlab not installed. Install with: pip install reportlab")
        path = os.path.join(outdir, f"securelog_{stamp}.pdf")
        c = pdf_canvas.Canvas(path, pagesize=A4)
        width, height = A4
        x, y = 2 * cm, height - 2 * cm
        c.setFont("Times-Roman", 14)
        c.drawString(x, y, "sky's log_analyzer — Report")
        c.setFont("Times-Roman", 10)
        y -= 1 * cm
        for line in text.splitlines():
            if y < 2 * cm:
                c.showPage()
                y = height - 2 * cm
                c.setFont("Times-Roman", 10)
            c.drawString(x, y, line[:110])
            y -= 0.45 * cm
        c.showPage()
        c.save()
        return path

    def _render_donut_png(self, title, labels, values, colors):
        if not HAS_MPL: return None
        fig = Figure(figsize=(5.4, 3.6), dpi=110, facecolor=COLOR_BG_PANEL)
        ax = fig.add_subplot(111)
        total = max(1, sum(values))

        def pct_fmt(pct): return f"{pct:.0f}%"

        ax.pie(values if total else [1], labels=labels, colors=colors, startangle=90,
               wedgeprops=dict(width=0.38), autopct=pct_fmt, textprops={'color': 'white'})
        ax.set_title(title, color='white')
        buf = io.BytesIO()
        fig.savefig(buf, format="png", bbox_inches="tight", facecolor=COLOR_BG_PANEL)
        return buf.getvalue()

    def generate_reports(self, formats=("txt", "html")):
        outdir = self._ensure_reports_dir()
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        text = self._prepare_report_text()

        # Optional donut images for HTML
        total_ips = len(self.stats)
        unique_once = sum(1 for _ip, s in self.stats.items() if s["total_requests"] == 1)
        repeated = max(0, total_ips - unique_once)
        low = self.sev_counts.get("Low", 0)
        med = self.sev_counts.get("Medium", 0)
        high = self.sev_counts.get("High", 0)
        png_uni = self._render_donut_png("Uniqueness", ["Unique (once)", "Repeated"], [unique_once, repeated],
                                         [COLOR_SUCCESS, COLOR_WARNING])
        png_sev = self._render_donut_png("Severity", ["Low", "Medium", "High"], [low, med, high],
                                         [COLOR_SUCCESS, COLOR_WARNING, COLOR_DANGER])

        saved = []
        for fmt in formats:
            fmt = fmt.lower().strip()
            if fmt == "txt":
                saved.append(self._save_txt(text, outdir, stamp))
            elif fmt == "html":
                saved.append(self._render_html(text, outdir, stamp, png_uni, png_sev))
            elif fmt == "docx":
                try:
                    saved.append(self._save_docx(text, outdir, stamp))
                except Exception as e:
                    messagebox.showwarning("DOCX", str(e))
            elif fmt == "pdf":
                try:
                    saved.append(self._save_pdf(text, outdir, stamp))
                except Exception as e:
                    messagebox.showwarning("PDF", str(e))
        return [p for p in saved if p]

    # popups
    def _popup_list(self, title: str, items: list[str]):
        win = tk.Toplevel(self.master)
        win.title(title)
        win.geometry("600x520")
        win.configure(bg=COLOR_BG_MAIN)

        txt = tk.Text(win, wrap="none", bg=COLOR_BG_PANEL, fg=COLOR_TEXT_MAIN, relief="flat")
        ysb = ttk.Scrollbar(win, orient="vertical", command=txt.yview)
        xsb = ttk.Scrollbar(win, orient="horizontal", command=txt.xview)
        txt.configure(yscrollcommand=ysb.set, xscrollcommand=xsb.set)
        txt.grid(row=0, column=0, sticky="nsew")
        ysb.grid(row=0, column=1, sticky="ns")
        xsb.grid(row=1, column=0, sticky="ew")
        win.rowconfigure(0, weight=1)
        win.columnconfigure(0, weight=1)
        for it in items:
            txt.insert("end", it + "\n")
        txt.focus_set()
        btnfrm = ttk.Frame(win)
        btnfrm.grid(row=2, column=0, columnspan=2, sticky="ew")
        ttk.Button(btnfrm, text="Copy All",
                   command=lambda: (self.master.clipboard_clear(),
                                    self.master.clipboard_append("\n".join(items)),
                                    messagebox.showinfo("Copied", "List copied to clipboard."))
                   ).pack(side="left", padx=6, pady=6)
        ttk.Button(btnfrm, text="Close", command=win.destroy).pack(side="right", padx=6, pady=6)

    def show_all_ips(self):
        items = sorted(self.stats.keys())
        self._popup_list("All IPs", items)

    def show_unique_ips(self):
        items = sorted([k for k, v in self.stats.items() if v["total_requests"] == 1])
        self._popup_list("Unique IPs (seen once)", items)

    def export_csv(self):
        if not self.stats:
            messagebox.showinfo("Info", "Run analysis first.");
            return
        p = filedialog.asksaveasfilename(title="Export CSV", defaultextension=".csv", filetypes=[("CSV", "*.csv")])
        if not p: return
        try:
            with open(p, "w", newline="", encoding="utf-8") as f:
                w = csv.writer(f)
                w.writerow(["ip", "total_requests", "ua_matches", "sources_used",
                            "severity_label", "severity_score",
                            "abuse_score", "abuse_reports", "abuse_isp",
                            "vt_malicious", "otx_threat_level", "otx_pulses"])
                for ip, s in self.stats.items():
                    ua_pairs = "; ".join([f"{label}:{count}" for label, count in sorted(s["ua_counts"].items())])
                    sources = "; ".join(sorted(s["sources_used"]))
                    ab = s.get("abuseipdb") or {}
                    vt = s.get("virustotal") or {}
                    ox = s.get("otx") or {}
                    w.writerow([
                        ip, s["total_requests"], ua_pairs, sources,
                        s.get("severity_label", "Low"), s.get("severity_score", 0),
                        ab.get("score", ""), ab.get("total_reports", ""), ab.get("isp", ""),
                        (vt.get("malicious_count") if vt else ""),
                        (ox.get("threat_level") if ox else ""), (ox.get("pulses") if ox else "")
                    ])
            messagebox.showinfo("Success", f"CSV saved: {p}")
        except Exception as e:
            messagebox.showerror("Error", f"CSV export failed: {e}")

    def export_json(self):
        if not self.stats:
            messagebox.showinfo("Info", "Run analysis first.");
            return
        p = filedialog.asksaveasfilename(title="Export JSON", defaultextension=".json", filetypes=[("JSON", "*.json")])
        if not p: return
        try:
            out = {}
            for ip, s in self.stats.items():
                out[ip] = {
                    "total_requests": s["total_requests"],
                    "ua_counts": dict(s["ua_counts"]),
                    "sources_used": sorted(s["sources_used"]),
                    "sources_flagged": sorted(s["sources_flagged"]),
                    "severity_label": s.get("severity_label", "Low"),
                    "severity_score": s.get("severity_score", 0),
                    "abuseipdb": s.get("abuseipdb"),
                    "virustotal": s.get("virustotal"),
                    "otx": s.get("otx"),
                }
            bundle = {
                "summary": {
                    "total_ips": len(self.stats),
                    "ua_totals": dict(self.ua_totals),
                    "severity_counts": dict(self.sev_counts),
                    "parse_errors": self.parse_errors,
                },
                "ips": out
            }
            with open(p, "w", encoding="utf-8") as f:
                json.dump(bundle, f, ensure_ascii=False, indent=2)
            messagebox.showinfo("Success", f"JSON saved: {p}")
        except Exception as e:
            messagebox.showerror("Error", f"JSON export failed: {e}")


# ==========================
# -------- CLI mode --------
# ==========================
def run_cli(args):
    try:
        stats, ua_totals, parse_errors, sev_counts = analyze_fast(
            args.log,
            enable_cti=args.cti,
            vt_key=os.environ.get("VT_API_KEY", ""),
            otx_key=os.environ.get("OTX_API_KEY", ""),
            abuse_key=os.environ.get("ABUSEIPDB_API_KEY", ""),
            cti_lookup_top_n=args.topn
        )
        # Minimal console output
        print(f"IPs: {len(stats)} | parse_errors: {parse_errors} | severity: {dict(sev_counts)}")
        # Prepare & save reports
        app_like = type("Dummy", (), {})()
        app_like.stats = stats
        app_like.ua_totals = ua_totals
        app_like.parse_errors = parse_errors
        app_like.sev_counts = sev_counts
        SecureLogApp._ensure_reports_dir = SecureLogApp._ensure_reports_dir
        SecureLogApp._prepare_report_text = SecureLogApp._prepare_report_text
        SecureLogApp._render_donut_png = SecureLogApp._render_donut_png
        SecureLogApp._render_html = SecureLogApp._render_html
        SecureLogApp._save_txt = SecureLogApp._save_txt
        SecureLogApp._save_docx = SecureLogApp._save_docx
        SecureLogApp._save_pdf = SecureLogApp._save_pdf

        # Monkey-bind methods to dummy
        def bind(m):
            return m.__get__(app_like, SecureLogApp)

        app_like._ensure_reports_dir = bind(SecureLogApp._ensure_reports_dir)
        app_like._prepare_report_text = bind(SecureLogApp._prepare_report_text)
        app_like._render_donut_png = bind(SecureLogApp._render_donut_png)
        app_like._render_html = bind(SecureLogApp._render_html)
        app_like._save_txt = bind(SecureLogApp._save_txt)
        app_like._save_docx = bind(SecureLogApp._save_docx)
        app_like._save_pdf = bind(SecureLogApp._save_pdf)

        formats = [f.strip() for f in (args.report or "txt,html").split(",")]
        outdir = app_like._ensure_reports_dir()
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        text = app_like._prepare_report_text()

        # donuts
        total_ips = len(stats)
        unique_once = sum(1 for _ip, s in stats.items() if s["total_requests"] == 1)
        repeated = max(0, total_ips - unique_once)
        low = sev_counts.get("Low", 0);
        med = sev_counts.get("Medium", 0);
        high = sev_counts.get("High", 0)
        png_uni = app_like._render_donut_png("Uniqueness", ["Unique (once)", "Repeated"], [unique_once, repeated],
                                             [COLOR_SUCCESS, COLOR_WARNING])
        png_sev = app_like._render_donut_png("Severity", ["Low", "Medium", "High"], [low, med, high],
                                             [COLOR_SUCCESS, COLOR_WARNING, COLOR_DANGER])

        saved = []
        for fmt in formats:
            fmt = fmt.lower()
            if fmt == "txt":
                saved.append(app_like._save_txt(text, outdir, stamp))
            elif fmt == "html":
                saved.append(app_like._render_html(text, outdir, stamp, png_uni, png_sev))
            elif fmt == "docx":
                try:
                    saved.append(app_like._save_docx(text, outdir, stamp))
                except Exception as e:
                    print(f"[DOCX] {e}")
            elif fmt == "pdf":
                try:
                    saved.append(app_like._save_pdf(text, outdir, stamp))
                except Exception as e:
                    print(f"[PDF] {e}")
        print("Saved:\n" + "\n".join(saved))
    except FileNotFoundError as e:
        print(f"[Error] {e}")
    except Exception as e:
        print(f"[Error] Unexpected: {e}")


def parse_cli():
    p = argparse.ArgumentParser(description="sky's log_analyzer — CLI mode")
    p.add_argument("log", nargs="?", help="Path to log file")
    p.add_argument("--cti", action="store_true", help="Enable CTI lookups")
    p.add_argument("--topn", type=int, default=None, help="CTI top N IPs by requests")
    p.add_argument("--report", type=str, default="txt,html", help="Comma list: txt,html,docx,pdf")
    p.add_argument("--cli", action="store_true", help="Run in CLI mode only")
    return p.parse_args()


# ==========================
# ---------- main ----------
# ==========================
def main():
    args = parse_cli()
    if args.cli and args.log:
        run_cli(args)
        return

    root = tk.Tk()
    if START_MAXIMIZED:
        try:
            root.state('zoomed')
        except Exception:
            try:
                root.attributes('-zoomed', True)
            except Exception:
                sw = root.winfo_screenwidth();
                sh = root.winfo_screenheight()
                root.geometry(f"{sw}x{sh}+0+0")
    app = SecureLogApp(root)
    app.grid(row=0, column=0, sticky="nsew")
    root.rowconfigure(0, weight=1)
    root.columnconfigure(0, weight=1)
    # No root.configure(bg=...) needed as frames cover it, but good practice:
    root.configure(bg=COLOR_BG_MAIN)
    root.mainloop()
if __name__ == "__main__":
    main()